#view_link_summaries.py
import os
import json
import streamlit as st
import logging
from datetime import datetime
from docx import Document
import re
from modules.helpers import load_config

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

config = load_config()

OUTPUT_FOLDER = config.get("file_paths", {}).get("output_article_summaries_word_docs") #"output"  #

def load_summary_files():
    summaries_dir = config['file_paths']['output_dir_article']
    logger.debug(f"Summaries directory path: {summaries_dir}")

    if not os.path.exists(summaries_dir):
        logger.error(f"The summaries directory does not exist: {summaries_dir}")
        return []

    json_files = [f for f in os.listdir(summaries_dir) if f.endswith('.json')]
    if not json_files:
        logger.error("No summary files found.")
        return []

    json_files.sort(key=lambda x: os.path.getmtime(os.path.join(summaries_dir, x)), reverse=True)
    file_options = [(f, datetime.fromtimestamp(os.path.getmtime(os.path.join(summaries_dir, f))).strftime('%Y-%m-%d %H:%M:%S')) for f in json_files]
    return file_options

def load_summary_content(file_name):
    summaries_dir = config['file_paths']['output_dir_article']
    file_path = os.path.join(summaries_dir, file_name)
    try:
        with open(file_path, 'r') as json_file:
            content = json.load(json_file)
        logger.debug(f"Successfully loaded content from {file_name}")
        return content
    except json.JSONDecodeError as e:
        logger.error(f"JSON decoding error in {file_name}: {e}")
        st.error(f"Error decoding JSON in {file_name}. Displaying raw content.")
        with open(file_path, 'r') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error loading content from {file_name}: {e}")
        return None

def display_dictionary_content(dictionary, selected_keys):
    """Display selected keys from a dictionary using Markdown."""
    for key in selected_keys:
        if key in dictionary:
            value = dictionary[key]
            st.markdown(f"#### {key}")
            if isinstance(value, list):
                list_items = '\n'.join(f"- {item}" for item in value)
                st.markdown(list_items)
            else:
                st.markdown(value)
            st.markdown("---")

def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def save_to_word(content: dict, json_file_name: str) -> str:
    """Save all parts of the displayed content to a Word file."""
    doc = create_document_with_heading(json_file_name)

    if isinstance(content, list):
        process_list_content(content, doc)
    else:
        add_non_list_content_to_doc(content, doc)

    return save_document(doc, json_file_name)

def create_document_with_heading(json_file_name: str) -> Document:
    """Create a new Document and add the main heading."""
    doc = Document()
    base_name = os.path.splitext(json_file_name)[0]
    doc.add_heading(base_name, level=1)
    return doc

def process_list_content(content: list, doc: Document):
    """Process content if it's a list of items."""
    for index, item in enumerate(content):
        if isinstance(item, dict):
            # Filtered item now uses all keys as they appear
            if all(is_na_value(value) for value in item.values()):
                continue  # Skip this article as it's considered 'N/A'

            add_dict_content_to_doc(item, doc, index)
        else:
            doc.add_paragraph(str(item))

def is_na_value(value):
    """Check if a value is considered 'N/A' or 'NA'."""
    if isinstance(value, list):
        return all(item.lower() in ('n/a', 'na', 'not available') for item in value if isinstance(item, str))
    if isinstance(value, str):
        return value.lower() in ('n/a', 'na', 'not available')
    return False

def add_dict_content_to_doc(item: dict, doc: Document, index: int):
    """Add dictionary content to the Word document."""
    if index > 0:
        doc.add_page_break()  # Add a page break before each new article
    doc.add_heading(f"Article {index + 1}", level=2)

    # Process each key-value pair in the item, preserving the original order of keys
    for key, value in item.items():
        add_key_value_to_doc(doc, key, value)

    # doc.add_paragraph("✨"*30)  # A visual separator between articles

def add_key_value_to_doc(doc: Document, key: str, value):
    """Add a key-value pair to the Word document."""
    clean_key = clean_key_text(key)
    doc.add_heading(clean_key, level=3)
    if isinstance(value, list):
        add_list_to_doc(value, doc)
    else:
        add_value_to_doc(key, value, doc)

def clean_key_text(key: str) -> str:
    """Clean the key text to make it more readable."""
    return key.replace('_', ' ').title().replace("\ud83d\udcbb", '').replace("\ud83d\udcf0", '')

def add_list_to_doc(value_list: list, doc: Document):
    """Add a list of values to the Word document as bullet points."""
    for list_item in value_list:
        doc.add_paragraph(f"- {list_item}")

def add_value_to_doc(key: str, value, doc: Document):
    """Add a value to the Word document, splitting long text if necessary."""
    if "summary" in key.lower() or "takeaway" in key.lower():
        sentences = value.split('. ')
        for sentence in sentences:
            doc.add_paragraph(f"{sentence.strip()}.")
    else:
        doc.add_paragraph(str(value))

def add_non_list_content_to_doc(content: dict, doc: Document):
    """Add non-list content to the Word document."""
    doc.add_paragraph(json.dumps(content, indent=4))

def save_document(doc: Document, json_file_name: str) -> str:
    """Save the Word document to the output folder."""
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    file_path = os.path.join(OUTPUT_FOLDER, f"{os.path.splitext(json_file_name)[0]}.docx")
    try:
        doc.save(file_path)
        logger.info(f"Saved {file_path} successfully!")
        st.session_state.docx_file_path = file_path
        return file_path
    except Exception as e:
        logger.error(f"Error saving Word file: {e}")
        return None


def display_download_button():
    if 'docx_file_path' in st.session_state:
        docx_file_path = st.session_state.docx_file_path
        try:
            with open(docx_file_path, "rb") as f:
                st.download_button(
                    label="Download .docx File",
                    data=f,
                    file_name=os.path.basename(docx_file_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except FileNotFoundError:
            st.error("The .docx file was not found. Please ensure it was saved correctly.")
        except Exception as e:
            st.error(f"Error opening the .docx file: {e}")

def view_link_summaries():
    file_options = load_summary_files()
    if not file_options:
        st.error("No summary files found.")
        return

    selected_file = st.selectbox("Select a summary file to view:", options=file_options, format_func=lambda x: f"{x[0]} (Last modified: {x[1]})")

    if selected_file:
        content = load_summary_content(selected_file[0])
        if content:
            process_selected_content(content, selected_file[0])
        else:
            st.error("Failed to load the selected file content.")

# def process_selected_content(content, json_file_name: str):
#     if isinstance(content, list):
#         item_options = [f"Article {i + 1}" for i in range(len(content))]
#         selected_items = st.multiselect("Select items to display:", item_options, default=item_options)
#         selected_indices = [int(item.split()[-1]) - 1 for item in selected_items]

#         possible_keys = collect_possible_keys(content, selected_indices)

#         selected_keys = st.multiselect(
#             "Select parts to display for selected items:",
#             possible_keys,
#             default=possible_keys
#         )

#         if st.checkbox("Display selected parts"):
#             for index in selected_indices:
#                 dictionary = content[index]
#                 if isinstance(dictionary, dict):
#                     display_dictionary_content(dictionary, selected_keys)
#                 else:
#                     st.error("The selected item is not a dictionary.")

#         if st.button("Save as .docx"):
#             if selected_keys:
#                 file_path = save_to_word(content, json_file_name, selected_keys)
#                 if file_path:
#                     display_download_button()
#             else:
#                 st.error("Please select parts to display before saving.")
#     else:
#         st.warning("Content is not in list format. Displaying all content as JSON.")
#         st.json(content)

def process_selected_content(content, json_file_name: str):
    if isinstance(content, list):
        # Automatically process all items in the content
        if st.checkbox("Show Summaries"):
            with st.container():
                st.markdown(f"""
                <div style="overflow-y: auto; height: 600px;">  <!-- Adjust the height as needed -->
                    {"".join(
                        f"<h3>Article {index + 1}</h3>" +
                        "".join(
                            f"<p><b>{key}:</b> {'<ul><li>' + '</li><li>'.join(map(str, value)) + '</li></ul>' if isinstance(value, list) else value}</p>"
                            for key, value in dictionary.items()
                        )
                        for index, dictionary in enumerate(content) if isinstance(dictionary, dict)
                    )}
                </div>
                """, unsafe_allow_html=True)

            # Handle non-dictionary items
            for index, item in enumerate(content):
                if not isinstance(item, dict):
                    st.error(f"Item {index + 1} is not a dictionary.")

        if st.button("Save as .docx"):
            file_path = save_to_word(content, json_file_name)
            if file_path:
                display_download_button()
            else:
                st.error("Error in saving the file.")
    else:
        st.warning("Content is not in list format. Displaying all content as JSON.")
        st.json(content)


def collect_possible_keys(content: list, selected_indices: list) -> list:
    """Collect all possible keys from the selected items in the content."""
    all_possible_keys = set()
    for index in selected_indices:
        if isinstance(content[index], dict):
            all_possible_keys.update(content[index].keys())
    return list(all_possible_keys)

if __name__ == "__main__":
    view_link_summaries()
