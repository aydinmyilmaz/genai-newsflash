# #youtube_summary_manager.py
# import streamlit as st
# import os
# import re
# import json
# from datetime import datetime
# import matplotlib.pyplot as plt

# from modules.youtube_data_handler import (
#     validate_youtube_link,
#     extract_video_ids,
#     get_video_transcript,
#     save_links,
#     load_links,
#     get_channel_id
# )

# from modules.helpers import load_config

# from docx import Document
# from docx.shared import Pt

# # Ensure the necessary directories exist
# if not os.path.exists('./sources'):
#     os.makedirs('./sources')
# if not os.path.exists('./data/transcripts'):
#     os.makedirs('./data/transcripts')

# config = load_config()
# # Set the path for the links file
# LINKS_FILE_PATH = youtube_links_file = config["file_paths"]["youtube_links_file"]

# def youtube_summary_manager():

#     # Initialize session state for links if not already done
#     if 'existing_links' not in st.session_state:
#         st.session_state.existing_links = load_links(youtube_links_file)

#     tabs = st.tabs(["Manage Links", "Retrieve Transcripts"])

#     with tabs[0]:
#         # Load links from session state
#         existing_links = st.session_state.existing_links

#         # Text area for entering links
#         link_input = st.text_area(
#             "Paste YouTube links here (one per line):",
#             value="\n".join(existing_links),
#             key='link_input'
#         )

#         # Validate links in real-time
#         input_links = [link.strip() for link in link_input.split('\n') if link.strip()]
#         channel_links = []
#         video_links = []
#         invalid_links = []

#         for link in input_links:
#             if validate_youtube_link(link):
#                 try:
#                     if "channel" in link or "user" in link or "@" in link:
#                         # Convert channel/user link or handle link to channel ID
#                         channel_id = get_channel_id(link)
#                         if channel_id:
#                             corrected_link = f"https://www.youtube.com/channel/{channel_id}"
#                             channel_links.append(corrected_link)
#                         else:
#                             st.warning(f"Failed to retrieve channel ID for link: {link}")
#                             invalid_links.append(link)
#                     elif "watch?v=" in link:
#                         video_id_match = re.search(r"v=([a-zA-Z0-9_-]+)", link)
#                         if video_id_match:
#                             video_id = video_id_match.group(1)
#                             if video_id:  # Check that a valid video ID was found
#                                 video_links.append(link)
#                             else:
#                                 st.warning(f"Could not extract video ID from link: {link}")
#                                 invalid_links.append(link)
#                         else:
#                             st.warning(f"Invalid video link format for link: {link}")
#                             invalid_links.append(link)
#                     else:
#                         st.warning(f"Link type not recognized: {link}")
#                         invalid_links.append(link)
#                 except Exception as e:
#                     st.error(f"Error processing link: {link}. Details: {str(e)}")
#                     invalid_links.append(link)
#             else:
#                 invalid_links.append(link)

#         if invalid_links:
#             st.error("The following links are invalid YouTube URLs:")
#             for link in invalid_links:
#                 st.write(f"- {link}")

#         # Buttons to Save/Update links and Display Saved Links
#         col1, col2 = st.columns(2)
#         with col1:
#             if st.button("💾 Update Links"):
#                 if invalid_links:
#                     st.error("Cannot save. Please fix invalid links.")
#                 else:
#                     # Save the validated links
#                     new_links = channel_links + video_links
#                     # st.write("Attempting to save the following links:", new_links)  # Debug statement
#                     try:
#                         save_links(new_links, youtube_links_file)  # Ensure this function is correctly saving the links
#                         st.session_state.existing_links = new_links
#                         st.success(f"Links saved successfully to {youtube_links_file}.")
#                     except Exception as e:
#                         st.error(f"Error saving links: {e}")

#                     # Reload the links from the file to ensure the latest data is displayed
#                     st.session_state.existing_links = load_links(youtube_links_file)  # Reload links after saving
#                     # st.write("Current links in session state:", st.session_state.existing_links)  # Debug statement

#         with col2:
#             if st.button("📜 Display Links"):
#                 # Reload the links from the session state
#                 st.write("Current links in session state:", st.session_state.existing_links)  # Debug statement


#     with tabs[1]:
#         # Number of videos to retrieve and Get Transcripts button
#         if channel_links or video_links:
#             col1, col2 = st.columns([1, 2])  # Adjust column widths for better alignment
#             with col1:
#                 num_videos = st.number_input("Select number of videos to retrieve from each channel:", min_value=1, max_value=10, value=2, step=1)
#             with col2:
#                 # Add explanatory text next to the button
#                 st.markdown("Click the button to retrieve transcripts of videos.")
#                 if st.button("Get Transcripts"):
#                     transcripts_data = {}
#                     token_counts = []
#                     video_titles = []
#                     video_links = []

#                     progress_bar = st.progress(0)
#                     total_links = len(channel_links)

#                     for idx, channel_id in enumerate(channel_links):
#                         st.write(f"Processing channel link {idx+1}/{len(channel_links)}: {channel_id}")

#                         video_ids = extract_video_ids(channel_id, max_results=num_videos)
#                         st.write(f"Video IDs: {video_ids}")
#                         for vid_idx, video_id in enumerate(video_ids):
#                             progress = (idx * num_videos + vid_idx + 1) / (total_links * num_videos)
#                             progress_bar.progress(min(progress, 1.0))

#                             video_link = f"https://www.youtube.com/watch?v={video_id}"
#                             transcript_data = get_video_transcript(video_id)
#                             if transcript_data:
#                                 # Only add to transcripts_data if token count is greater than 1000
#                                 if transcript_data['token_count'] > 1000:
#                                     transcripts_data[video_id] = {
#                                         'transcript': transcript_data['transcript'],
#                                         'token_count': transcript_data['token_count'],
#                                         'link': video_link
#                                     }
#                                     token_counts.append(transcript_data['token_count'])
#                                     video_titles.append(f"Video {vid_idx+1} (from channel link {idx+1})")
#                                     video_links.append(video_link)
#                                 else:
#                                     st.warning(f"Transcript for video {video_id} has a token count of {transcript_data['token_count']}, which is less than 1000 and will not be saved.")
#                             else:
#                                 transcripts_data[video_id] = {
#                                     'transcript': None,
#                                     'token_count': 0,
#                                     'link': video_link
#                                 }
#                                 token_counts.append(0)
#                                 video_titles.append(f"Video {vid_idx+1} (from channel link {idx+1})")
#                                 video_links.append(video_link)

#                     # Store transcripts data in session state to persist between reruns
#                     st.session_state.transcripts_data = transcripts_data
#                     st.session_state.video_titles = video_titles

#                     # Save transcripts data to JSON only if there are valid entries
#                     if transcripts_data:
#                         current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
#                         output_file_path = f"{config['file_paths']['youtube_transcripts_folder']}/transcripts_{current_time}.json"
#                         with open(output_file_path, 'w', encoding='utf-8') as f:
#                             json.dump(transcripts_data, f, ensure_ascii=False, indent=4)

#                         st.success(f"Transcript retrieval completed. Results saved to '{output_file_path}'.")
#                     else:
#                         st.warning("No transcripts with token counts greater than 1000 were found.")

#         # Display Transcripts using a Dropdown Selection (fixed expander issue)
#         if "transcripts_data" in st.session_state and st.session_state.transcripts_data and st.checkbox("Display Transcript"):
#             dropdown_options = [
#                 f"{st.session_state.video_titles[idx]} - Preview: {data['transcript'][:100] if data['transcript'] else 'Transcript not available'}"
#                 for idx, data in enumerate(st.session_state.transcripts_data.values())
#             ]

#             selected_option = st.selectbox("Select a transcript to view:", dropdown_options)
#             selected_idx = dropdown_options.index(selected_option)
#             selected_video_id = list(st.session_state.transcripts_data.keys())[selected_idx]
#             selected_data = st.session_state.transcripts_data[selected_video_id]

#             st.subheader(f"{st.session_state.video_titles[selected_idx]} - Token Count: {selected_data['token_count']}")
#             st.write(f"Video Link: {selected_data['link']}")
#             if selected_data['transcript']:
#                 st.write(selected_data['transcript'][:1000])
#             else:
#                 st.warning("Transcript not available.")

# def sanitize_filename(filename):
#     return re.sub(r'[<>:"/\\|?*]', '_', filename)

# def save_youtube_summary_to_docx(content, json_file_name):
#     """Save all items from the content with formatted text in an improved style to a .docx file."""
#     doc = Document()
#     base_name = os.path.splitext(json_file_name)[0]  # Get the base name without extension
#     doc.add_heading(base_name, level=1)

#     def add_heading(doc, text, level):
#         """Helper function to add headings in DOCX."""
#         doc.add_heading(text, level=level)

#     def add_bullet_points(doc, items):
#         """Helper function to add bullet points in DOCX."""
#         for item in items:
#             doc.add_paragraph(item.strip('- '), style='List Bullet')

#     def add_paragraph(doc, text):
#         """Helper function to add paragraphs with Markdown-style formatting converted."""
#         paragraph = doc.add_paragraph()

#         # Split the text using regex to identify **bold**, _italic_, or both combined
#         elements = re.split(r'(\*\*.+?\*\*|_.+?_)', text)
#         for element in elements:
#             if element.startswith('**') and element.endswith('**'):
#                 run = paragraph.add_run(element.strip('**'))
#                 run.bold = True
#             elif element.startswith('_') and element.endswith('_'):
#                 run = paragraph.add_run(element.strip('_'))
#                 run.italic = True
#             else:
#                 paragraph.add_run(element)

#         paragraph.style.font.size = Pt(11)

#     def add_separator(doc):
#         """Helper function to add a horizontal line separator between items."""
#         paragraph = doc.add_paragraph()
#         run = paragraph.add_run()
#         run.add_break()
#         run.add_text("✨" * 30)
#         run.add_break()

#     if isinstance(content, dict):
#         for video_id, item in content.items():
#             if isinstance(item, dict):
#                 for key, value in item.items():
#                     if key.lower() not in ['token_count', 'link']:
#                         if isinstance(value, str):
#                             lines = value.split('\n')
#                             for line in lines:
#                                 line = line.strip()
#                                 if line.startswith('###'):
#                                     add_heading(doc, line.strip('# '), level=2)
#                                 elif line.startswith('-'):
#                                     doc.add_paragraph(line.strip('- '), style='List Bullet')
#                                 elif line:
#                                     add_paragraph(doc, line)
#                         elif isinstance(value, list):
#                             add_bullet_points(doc, value)
#                         else:
#                             add_paragraph(doc, str(value))

#                 # Add link at the end of each video summary
#                 if 'link' in item:
#                     doc.add_paragraph(f"Video Link: {item['link']}")

#                 add_separator(doc)  # Add a separator after each item
#             else:
#                 add_paragraph(doc, f"{key}: {str(item)}")
#                 add_separator(doc)

#     output_folder = "output/youtube_summaries"
#     os.makedirs(output_folder, exist_ok=True)
#     file_path = os.path.join(output_folder, f"{sanitize_filename(base_name)}.docx")
#     try:
#         doc.save(file_path)
#         st.success(f"Saved {file_path} successfully!")
#         return file_path  # Return the file path for download
#     except Exception as e:
#         logger.error(f"Error saving .docx file: {e}")
#         st.error("Failed to save the .docx file.")
#         return None

# if __name__ == "__main__":
#     youtube_summary_manager()


import os
import re
import json
import logging
from datetime import datetime
from typing import List, Dict, Optional

import streamlit as st
from docx import Document
from docx.shared import Pt
import matplotlib.pyplot as plt

from modules.youtube_data_handler import (
    validate_youtube_link,
    extract_video_ids,
    get_video_transcript,
    save_links,
    load_links,
    get_channel_id
)
from modules.helpers import load_config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure the necessary directories exist
os.makedirs('./sources', exist_ok=True)
os.makedirs('./data/transcripts', exist_ok=True)

config = load_config()
LINKS_FILE_PATH = config["file_paths"]["youtube_links_file"]


def youtube_summary_manager() -> None:
    """
    Manages the Streamlit UI for handling YouTube links and retrieving transcripts.
    """
    # Initialize session state for links
    if 'existing_links' not in st.session_state:
        st.session_state.existing_links = load_links(LINKS_FILE_PATH)

    tabs = st.tabs(["Manage Links", "Retrieve Transcripts"])

    with tabs[0]:
        manage_links_tab()
    with tabs[1]:
        retrieve_transcripts_tab()


def manage_links_tab() -> None:
    """
    Handles the UI logic for managing YouTube links.
    """
    existing_links = st.session_state.existing_links

    link_input = st.text_area(
        "Paste YouTube links here (one per line):",
        value="\n".join(existing_links),
        key='link_input'
    )

    input_links = [link.strip() for link in link_input.split('\n') if link.strip()]
    channel_links, video_links, invalid_links = validate_links(input_links)

    if invalid_links:
        st.error("The following links are invalid YouTube URLs:")
        for link in invalid_links:
            st.write(f"- {link}")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Update Links"):
            update_links(channel_links, video_links, invalid_links)
    with col2:
        if st.button("📜 Display Links"):
            st.write("Current links in session state:", st.session_state.existing_links)


def retrieve_transcripts_tab() -> None:
    """
    Handles the UI logic for retrieving video transcripts.
    """
    channel_links = [link for link in st.session_state.existing_links if "channel" in link]
    video_links = [link for link in st.session_state.existing_links if "watch?v=" in link]

    if channel_links or video_links:
        col1, col2 = st.columns([1, 2])
        with col1:
            num_videos = st.number_input("Select number of videos to retrieve from each channel:", min_value=1, max_value=10, value=2, step=1)
        with col2:
            st.markdown("Click the button to retrieve transcripts of videos.")
            if st.button("Get Transcripts"):
                retrieve_and_save_transcripts(channel_links, video_links, num_videos)

    if "transcripts_data" in st.session_state and st.session_state.transcripts_data and st.checkbox("Display Transcript"):
        display_transcript_dropdown()


def validate_links(input_links: List[str]) -> (List[str], List[str], List[str]):
    """
    Validate YouTube links and categorize them as channel, video, or invalid links.
    """
    channel_links, video_links, invalid_links = [], [], []
    for link in input_links:
        if validate_youtube_link(link):
            try:
                if "channel" in link or "user" in link or "@" in link:
                    channel_id = get_channel_id(link)
                    if channel_id:
                        corrected_link = f"https://www.youtube.com/channel/{channel_id}"
                        channel_links.append(corrected_link)
                    else:
                        logger.warning(f"Failed to retrieve channel ID for link: {link}")
                        invalid_links.append(link)
                elif "watch?v=" in link:
                    video_id_match = re.search(r"v=([a-zA-Z0-9_-]+)", link)
                    if video_id_match:
                        video_links.append(link)
                    else:
                        logger.warning(f"Invalid video link format for link: {link}")
                        invalid_links.append(link)
                else:
                    logger.warning(f"Link type not recognized: {link}")
                    invalid_links.append(link)
            except Exception as e:
                logger.error(f"Error processing link: {link}. Details: {str(e)}")
                invalid_links.append(link)
        else:
            invalid_links.append(link)
    return channel_links, video_links, invalid_links


def update_links(channel_links: List[str], video_links: List[str], invalid_links: List[str]) -> None:
    """
    Updates the links by saving valid links and displaying appropriate error messages.
    """
    if invalid_links:
        st.error("Cannot save. Please fix invalid links.")
    else:
        new_links = channel_links + video_links
        try:
            save_links(new_links, LINKS_FILE_PATH)
            st.session_state.existing_links = new_links
            st.success(f"Links saved successfully to {LINKS_FILE_PATH}.")
        except Exception as e:
            logger.error(f"Error saving links: {e}")
            st.error(f"Error saving links: {e}")


def retrieve_and_save_transcripts(channel_links: List[str], video_links: List[str], num_videos: int) -> None:
    """
    Retrieves transcripts for videos from the given channel and video links and saves them.
    """
    transcripts_data = {}
    progress_bar = st.progress(0)
    total_links = len(channel_links) + len(video_links)

    for idx, channel_id in enumerate(channel_links):
        st.write(f"Processing channel link {idx+1}/{len(channel_links)}: {channel_id}")

        try:
            video_ids = extract_video_ids(channel_id, max_results=num_videos)
            for vid_idx, video_id in enumerate(video_ids):
                progress = (idx * num_videos + vid_idx + 1) / (total_links * num_videos)
                progress_bar.progress(min(progress, 1.0))

                video_link = f"https://www.youtube.com/watch?v={video_id}"
                transcript_data = get_video_transcript(video_id)
                if transcript_data and transcript_data['token_count'] > 1000:
                    transcripts_data[video_id] = {
                        'transcript': transcript_data['transcript'],
                        'token_count': transcript_data['token_count'],
                        'link': video_link
                    }
                else:
                    logger.warning(f"Transcript for video {video_id} has insufficient token count or is unavailable.")
        except Exception as e:
            logger.error(f"Error retrieving transcripts for channel {channel_id}: {e}")

    for idx, video_link in enumerate(video_links):
        st.write(f"Processing video link {idx+1}/{len(video_links)}: {video_link}")

        try:
            video_id_match = re.search(r"v=([a-zA-Z0-9_-]+)", video_link)
            if video_id_match:
                video_id = video_id_match.group(1)
                transcript_data = get_video_transcript(video_id)
                if transcript_data and transcript_data['token_count'] > 1000:
                    transcripts_data[video_id] = {
                        'transcript': transcript_data['transcript'],
                        'token_count': transcript_data['token_count'],
                        'link': video_link
                    }
                else:
                    logger.warning(f"Transcript for video {video_id} has insufficient token count or is unavailable.")
        except Exception as e:
            logger.error(f"Error retrieving transcript for video {video_link}: {e}")

    if transcripts_data:
        save_transcripts(transcripts_data)
    else:
        st.warning("No transcripts with token counts greater than 1000 were found.")


def save_transcripts(transcripts_data: Dict) -> None:
    """
    Saves the retrieved transcripts to a JSON file.
    """
    current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file_path = f"{config['file_paths']['youtube_transcripts_folder']}/transcripts_{current_time}.json"
    try:
        with open(output_file_path, 'w', encoding='utf-8') as f:
            json.dump(transcripts_data, f, ensure_ascii=False, indent=4)
        st.success(f"Transcript retrieval completed. Results saved to '{output_file_path}'.")
    except Exception as e:
        logger.error(f"Error saving transcripts to JSON: {e}")
        st.error("Failed to save transcripts to JSON.")


def display_transcript_dropdown() -> None:
    """
    Displays a dropdown to select and view transcripts.
    """
    dropdown_options = [
        f"{st.session_state.video_titles[idx]} - Preview: {data['transcript'][:100] if data['transcript'] else 'Transcript not available'}"
        for idx, data in enumerate(st.session_state.transcripts_data.values())
    ]

    selected_option = st.selectbox("Select a transcript to view:", dropdown_options)
    selected_idx = dropdown_options.index(selected_option)
    selected_video_id = list(st.session_state.transcripts_data.keys())[selected_idx]
    selected_data = st.session_state.transcripts_data[selected_video_id]

    st.subheader(f"{st.session_state.video_titles[selected_idx]} - Token Count: {selected_data['token_count']}")
    st.write(f"Video Link: {selected_data['link']}")
    if selected_data['transcript']:
        st.write(selected_data['transcript'][:1000])
    else:
        st.warning("Transcript not available.")


def sanitize_filename(filename: str) -> str:
    """
    Sanitizes the filename by removing or replacing invalid characters.
    """
    return re.sub(r'[<>:"/\\|?*]', '_', filename)


def save_youtube_summary_to_docx(content: Dict, json_file_name: str) -> Optional[str]:
    """
    Saves the YouTube summary content to a .docx file.

    Args:
        content (Dict): The content to save.
        json_file_name (str): The name of the JSON file.

    Returns:
        Optional[str]: The file path of the saved document, or None if an error occurred.
    """
    doc = Document()
    base_name = os.path.splitext(json_file_name)[0]
    doc.add_heading(base_name, level=1)

    try:
        for video_id, item in content.items():
            add_content_to_doc(doc, item)
            if 'link' in item:
                doc.add_paragraph(f"Video Link: {item['link']}")
            add_separator(doc)

        output_folder = "output/youtube_summaries"
        os.makedirs(output_folder, exist_ok=True)
        file_path = os.path.join(output_folder, f"{sanitize_filename(base_name)}.docx")
        doc.save(file_path)
        st.success(f"Saved {file_path} successfully!")
        return file_path
    except Exception as e:
        logger.error(f"Error saving .docx file: {e}")
        st.error("Failed to save the .docx file.")
        return None


def add_content_to_doc(doc: Document, item: Dict) -> None:
    """
    Adds the provided content to the Word document.

    Args:
        doc (Document): The Word document object.
        item (Dict): The content to add.
    """
    for key, value in item.items():
        if key.lower() not in ['token_count', 'link']:
            if isinstance(value, str):
                lines = value.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('###'):
                        add_heading(doc, line.strip('# '), level=2)
                    elif line.startswith('-'):
                        doc.add_paragraph(line.strip('- '), style='List Bullet')
                    elif line:
                        add_paragraph(doc, line)
            elif isinstance(value, list):
                add_bullet_points(doc, value)
            else:
                add_paragraph(doc, str(value))


def add_heading(doc: Document, text: str, level: int) -> None:
    """
    Adds a heading to the Word document.

    Args:
        doc (Document): The Word document object.
        text (str): The heading text.
        level (int): The heading level.
    """
    doc.add_heading(text, level=level)


def add_bullet_points(doc: Document, items: List[str]) -> None:
    """
    Adds bullet points to the Word document.

    Args:
        doc (Document): The Word document object.
        items (List[str]): The bullet points to add.
    """
    for item in items:
        doc.add_paragraph(item.strip('- '), style='List Bullet')


def add_paragraph(doc: Document, text: str) -> None:
    """
    Adds a paragraph to the Word document, with support for Markdown-style formatting.

    Args:
        doc (Document): The Word document object.
        text (str): The paragraph text.
    """
    paragraph = doc.add_paragraph()
    elements = re.split(r'(\*\*.+?\*\*|_.+?_)', text)
    for element in elements:
        if element.startswith('**') and element.endswith('**'):
            run = paragraph.add_run(element.strip('**'))
            run.bold = True
        elif element.startswith('_') and element.endswith('_'):
            run = paragraph.add_run(element.strip('_'))
            run.italic = True
        else:
            paragraph.add_run(element)
    paragraph.style.font.size = Pt(11)


def add_separator(doc: Document) -> None:
    """
    Adds a separator line to the Word document.

    Args:
        doc (Document): The Word document object.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break()
    run.add_text("\u2728" * 30)
    run.add_break()


if __name__ == "__main__":
    youtube_summary_manager()
