#smart_link_sumarizer.py
import streamlit as st
import json
import os

import tempfile
import logging
from datetime import datetime
import time
import hashlib
from docx import Document # Import the search function

from modules.searcher.search_filter import SearchModule
from modules.mongodb_manager.mongo_summary_handler import SummaryMongoHandler
from modules.utils.helpers import auth_check

from modules.scraper.smart_graph_scraper import smart_graph_scrap, load_config

from modules.utils.link_content_manager import link_content_manager
from modules.utils.article_summary_exporter import view_export_article_summaries
# Set page configuration at the top
st.set_page_config(layout="wide", page_title="Smart Link Summarizer")

# Set up logging
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(module)s - %(funcName)s - %(lineno)d - %(message)s')
logger = logging.getLogger(__name__)

# Add custom CSS for styling - moved near the top for cleaner organization
st.markdown(
    """
    <style>
    .stButton > button {
        background-color: #007BFF;
        color: white;
        border: 2px solid #007BFF;
        padding: 10px 20px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        margin: 4px 2px;
        cursor: pointer;
        border-radius: 5px;
        transition: background-color 0.3s, border-color 0.3s, color 0.3s;
    }
    .stButton > button:hover {
        background-color: #0056b3;
        border-color: #0056b3;
        color: white;
    }
    .stButton > button:focus {
        outline: none;
        box-shadow: 0 0 0 2px rgba(0, 123, 255, 0.5);
    }
    .stTextInput input, .stTextArea textarea {
        border: 2px solid #007BFF;
        border-radius: 5px;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: #0056b3;
    }
    .stCheckbox > label {
        color: #007BFF;
    }
    .stMarkdown {
        color: #333;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Authentication check at the start
auth_check()
# Get current user's email from session state
user_email = st.session_state.user_email

# Set up logging
logging.basicConfig(level=logging.INFO,
                   format='%(asctime)s - %(levelname)s - %(module)s - %(funcName)s - %(lineno)d - %(message)s')
logger = logging.getLogger(__name__)


def tooltip(text, help_text):
    """
    Returns a tooltip HTML string for the given text and help text.

    Args:
        text (str): The text to display.
        help_text (str): The help text for the tooltip.

    Returns:
        str: The HTML string with the tooltip.
    """
    return f'<span title="{help_text}">{text}</span>'


def display_links_selection(config):
    """
    Displays the list of available links to select from.

    Args:
        config (dict): The configuration dictionary.

    Returns:
        list: List of URLs available for processing.
    """
    try:
        with open(config['file_paths']['edited_link_file'], 'r') as file:
            urls = file.read().splitlines()
    except Exception as e:
        st.error(f"Error loading link data: {e}")
        logger.error(f"Error loading link data: {e}", exc_info=True)
        return []

    return urls


def save_results_to_file(results, config):
    """
    Saves the processed results to a JSON file with a timestamp.

    Args:
        results (list): The list of processed results.
        config (dict): The configuration dictionary.

    Returns:
        str: The file path of the saved results.
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"article_summary_results_{timestamp}.json"

    # Get the output directory from config, or use a default if not found
    output_dir = config.get('file_paths', {}).get('output_dir_article')

    file_path = os.path.join(output_dir, filename)

    os.makedirs(output_dir, exist_ok=True)

    # Save the results as a list of dictionaries
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=4, ensure_ascii=False)

    return file_path


def save_results_to_docx(results):
    """
    Saves the processed results to a DOCX file with a timestamp.

    Args:
        results (list): The list of processed results.

    Returns:
        str: The file path of the saved DOCX file.
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"article_summary_results_{timestamp}.docx"
    file_path = os.path.join(tempfile.gettempdir(), filename)

    doc = Document()
    doc.add_heading('Article Summary Results', level=1)

    for result in results:
        doc.add_heading(result.get('url', 'URL not available'), level=2)
        for key, value in result.items():
            # if key != 'url':
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    doc.save(file_path)
    return file_path


def cached_process_url(url, config, prompt):
    """
    Processes a URL using a cached function to avoid repeated processing of the same URL.
    """
    logger.info(f"Starting to process URL: {url}")

    try:
        logger.debug(f"Calling smart_graph_scrap with URL: {url}")
        result = smart_graph_scrap([url], config_path='config.yml', custom_prompt=prompt)
        logger.debug(f"Result from smart_graph_scrap: {result}")
        logger.debug(f"Type of result: {type(result)}")

        if isinstance(result, list):
            logger.info("Result is a list")
            if result:
                first_item = result[0]
                if isinstance(first_item, dict):
                    logger.info(f"Returning first item of list: {first_item}")
                    return first_item
                else:
                    logger.warning(f"First item is not a dictionary: {first_item}")
                    return {"url": url, "error": "Unexpected result structure"}
            else:
                logger.warning("Result list is empty")
                return {"url": url, "error": "Empty result list"}
        elif isinstance(result, dict):
            logger.info(f"Result is a dictionary: {result}")
            return result
        else:
            logger.warning(f"Unexpected result type: {type(result)}")
            return {"url": url, "error": f"Unexpected result type: {type(result)}"}
    except Exception as e:
        logger.error(f"Error processing URL {url}: {str(e)}", exc_info=True)
        return {"url": url, "error": str(e)}
    finally:
        logger.info(f"Finished processing URL: {url}")


def process_urls(urls, config, prompt):
    """
    Processes a list of URLs and returns the results.
    """
    results = []
    skipped_results = []  # Add this to track skipped articles
    start_time = time.time()

    # Create a progress bar
    progress_bar = st.progress(0)
    status_text = st.empty()

    # Initialize MongoDB handler
    mongo_handler = SummaryMongoHandler()

    # Filter out already processed URLs
    urls_to_process = []
    skipped_urls = []

    for url in urls:
        if mongo_handler.check_article_exists(url):
            skipped_urls.append(url)
            # Add skipped article to results with a skip flag
            skipped_results.append({
                "url": url,
                "skipped": True  # Add flag to indicate this was skipped
            })
        else:
            urls_to_process.append(url)

    # Show status of skipped URLs
    if skipped_urls:
        st.info(f"Skipping {len(skipped_urls)} already processed articles")
        logger.info(f"Skipping {len(skipped_urls)} already processed articles")

    # If no URLs to process, return early with skipped results
    if not urls_to_process:
        end_time = time.time()
        total_duration = end_time - start_time
        status_text.text("No new articles to process")
        return skipped_results, total_duration  # Return skipped results

    # Process remaining URLs
    for i, url in enumerate(urls_to_process, 1):
        url_start_time = time.time()
        status_text.text(f"Processing URL {i}/{len(urls_to_process)}: {url}")

        result = cached_process_url(url, config, prompt)
        if isinstance(result, dict):
            results.append(result)

        url_end_time = time.time()
        url_duration = url_end_time - url_start_time
        logger.info(f"Finished processing URL {i}/{len(urls_to_process)} in {url_duration:.2f} seconds")

        # Update progress bar
        progress_bar.progress(i / len(urls_to_process))

    end_time = time.time()
    total_duration = end_time - start_time
    logger.info(f"Total processing time: {total_duration:.2f} seconds")
    status_text.text(f"Processing completed in {total_duration:.2f} seconds")

    # Combine processed and skipped results
    all_results = results + skipped_results
    return all_results, total_duration


def load_prompts(prompt_folder):
    """
    Loads prompts from the specified folder.

    Args:
        prompt_folder (str): The folder containing the prompt files.

    Returns:
        dict: Dictionary with prompt filenames as keys and prompt text as values.
    """
    prompts = {}
    for filename in os.listdir(prompt_folder):
        if filename.endswith('.txt'):
            with open(os.path.join(prompt_folder, filename), 'r') as f:
                prompts[filename] = f.read().strip()
    return prompts


def load_main_prompt(config):
    """
    Loads the main prompt from the prompt folder.

    Args:
        config (dict): The configuration dictionary.

    Returns:
        str: The main prompt text, if available.
    """
    prompt_folder = config['file_paths']['prompt_folder']
    main_prompt_path = os.path.join(prompt_folder, 'main_prompt.txt')

    if os.path.exists(main_prompt_path):
        with open(main_prompt_path, 'r') as f:
            return f.read().strip()
    else:
        return None


def load_and_edit_prompt(config):
    """
    Main function to load, edit, and save prompts.

    Args:
        config (dict): The configuration dictionary.

    Returns:
        str: The edited prompt text.
    """
    prompt_folder = config['file_paths']['prompt_folder']

    # Step 1: Load available prompts and main prompt
    prompts, main_prompt = load_prompts_and_main(prompt_folder, config)

    if not prompts and main_prompt is None:
        st.warning("No prompt files found and no main_prompt.txt exists.")
        return None

    # Step 2: Select a prompt to edit
    selected_prompt, current_prompt = select_and_display_prompt(prompts, main_prompt)
    if current_prompt is None:
        return None

    # Step 3: Edit and save the selected prompt
    edited_prompt = edit_prompt(current_prompt)

    # Step 4: Save the edited prompt to the appropriate files
    if st.button("Save Prompt"):
        save_edited_prompt(prompt_folder, selected_prompt, edited_prompt)

    return edited_prompt


def load_prompts_and_main(prompt_folder, config):
    """Loads available prompts and the main prompt."""
    prompts = load_prompts(prompt_folder)
    main_prompt = load_main_prompt(config)
    return prompts, main_prompt


def select_and_display_prompt(prompts, main_prompt):
    """Handles prompt selection and displays the selected prompt."""
    prompt_options = list(prompts.keys()) + ["main_prompt.txt"]
    selected_prompt = st.selectbox("Select a prompt", prompt_options)

    if selected_prompt == "main_prompt.txt":
        current_prompt = main_prompt if main_prompt is not None else ""
    elif selected_prompt:
        current_prompt = prompts[selected_prompt]
    else:
        st.warning("Please select a prompt.")
        current_prompt = None

    return selected_prompt, current_prompt


def edit_prompt(current_prompt):
    """Displays a text area for the user to edit the prompt."""
    return st.text_area("Edit Prompt", current_prompt, height=300)


def save_edited_prompt(prompt_folder, selected_prompt, edited_prompt):
    """Saves the edited prompt to main_prompt.txt and, optionally, another file."""
    try:
        # Save the edited prompt to main_prompt.txt
        main_prompt_path = os.path.join(prompt_folder, 'main_prompt_link.txt')
        save_prompt_to_file(main_prompt_path, edited_prompt)
        st.success("Prompt saved successfully to main_prompt.txt!")

        # Save the edited prompt to the selected prompt file (if not main_prompt.txt)
        if selected_prompt != "main_prompt.txt":
            prompt_path = os.path.join(prompt_folder, selected_prompt)
            save_prompt_to_file(prompt_path, edited_prompt)
            st.success(f"Prompt also saved to {selected_prompt}!")
    except Exception as e:
        st.error(f"Error saving prompt: {e}")


def save_prompt_to_file(file_path, content):
    """Saves the given content to the specified file path."""
    with open(file_path, 'w') as f:
        f.write(content)


def display_article_details(results):
    """
    Displays the details of a selected article.

    Args:
        results (dict): Dictionary of processed articles.
    """
    if 'results' not in st.session_state or not st.session_state.results:
        st.warning("No articles loaded to display. Please process links first.")
        return

    article_ids = list(range(len(st.session_state.results)))  # Assuming each result corresponds to an article
    selected_id = st.selectbox("Select an article number to display:", article_ids)

    if selected_id is not None and selected_id < len(st.session_state.results):
        article = st.session_state.results[selected_id]
        if article:
            with st.expander("Article Details", expanded=True):
                for key, value in article.items():
                    if isinstance(value, list):
                        st.write(f"**{key.replace('_', ' ').title()}:**")
                        for idx, item in enumerate(value):
                            if isinstance(item, dict):
                                st.write(f"  - **{item.get('author', 'Author Unknown')}** says: \"{item.get('quote', 'No quote available')}\"")
                            else:
                                st.write(f"- {item}")
                    elif isinstance(value, dict):
                        for subkey, subvalue in value.items():
                            st.write(f"**{subkey.replace('_', ' ').title()}:** {subvalue}")
                    else:
                        st.write(f"**{key.replace('_', ' ').title()}:** {value}")
        else:
            st.error("Selected article data is not available.")


def is_valid_url(url):
    # List of file extensions to exclude
    excluded_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.webp', '.ico', '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx']

    # Check if the URL ends with any of the excluded extensions
    return not any(url.lower().endswith(ext) for ext in excluded_extensions)


def fetch_ai_news_links():
    """Main function to fetch AI news links and manage the UI flow."""
    config_2 = load_config()

    # Step 1: User Input for URLs and Number of Links
    urls, n_value = get_user_input()

    # Step 2: Handle URL processing if the "Generate Links" button is clicked
    if st.button("Generate Links"):
        if urls:
            process_urls_and_save_results(urls, n_value, config_2)
        else:
            st.warning("Please enter at least one URL.")


def get_user_input():
    """Handles user input for URLs and number of links to retrieve."""
    urls_input = st.text_area("Enter URLs (one per line)", height=200, help="Input up to 10 URLs to be processed.")

    n_value_portal = st.number_input(
        "Select number of portals to process:",
        min_value=1,
        max_value=50,
        value=10,
        help="Set how many links to retrieve from each URL."
    )
    urls = urls_input.strip().splitlines()[:n_value_portal]  # Limit input to 10 URLs

    n_value = st.number_input(
        "Select number of links to retrieve from each URL:",
        min_value=1,
        max_value=20,
        value=5,
        help="Set how many links to retrieve from each URL."
    )

    return urls, n_value


def process_urls_and_save_results(urls, n_value, config_2):
    """Processes URLs and saves the results to JSON and text files."""
    with st.spinner("Processing URLs..."):
        combined_results, successful_count = fetch_and_process_links(urls, n_value)

        if combined_results:
            # Step 3: Save the combined results
            save_combined_results(combined_results, successful_count, config_2)
        else:
            st.warning("No links found to save.")


def fetch_and_process_links(urls, n_value):
    """Fetches and processes links from the provided URLs."""
    prompt = (
        "Return full url (to access the article), title and date (usually found at the top of article) of available articles from the extracted web page data as json. "
        "Name main json key as portal name. The JSON should be in the following format: "
        "{\"WIRED\": [\n"
        "    {\"url\": \"https://www.wired.com/story/worldcoin-sam-altman-orb/\", \"title\": \"Sam Altman’s Eye-Scanning Orb Has a New Look—and Will Come Right to Your Door\", \"datetime\": \"May 31, 2023\"},\n"
        "    {\"url\": \"https://www.wired.com/story/filmmakers-are-worried-about-ai-big-tech-wants-them-to-see-whats-possible/\", \"title\": \"Filmmakers Are Worried About AI. Big Tech Wants Them to See 'What's Possible'\", \"datetime\": \"1 day ago\"}\n"
        "]}"
    )

    combined_results = {}
    successful_count = 0

    latest_success_message = None
    for i, url in enumerate(urls):
        try:
            result = process_single_url(url, prompt, n_value)
            if result:
                # Merge the result into combined_results
                combined_results = merge_results(combined_results, result, n_value)
                latest_success_message = f"Links from URL {i + 1} processed successfully."
                successful_count += 1
            else:
                st.warning(f"No links found for URL {i + 1}: {url}")
        except Exception as e:
            st.error(f"Error processing URL {i + 1}: {url}. Error: {e}")

    if latest_success_message:
        st.success(latest_success_message)

    return combined_results, successful_count


def process_single_url(url, prompt, n_value):
    """Processes a single URL and returns the result."""
    try:
        result, _ = process_urls([url], load_config(), prompt)
        if isinstance(result, str) and result.startswith("Links:"):
            result = result.replace("Links:", "").strip()
            result = json.loads(result)  # Convert string to JSON

        return result
    except Exception as e:
        logger.error(f"Error processing URL {url}: {e}", exc_info=True)
        return None


def merge_results(combined_results, result, n_value):
    """Merges the result from a single URL into the combined results."""
    if isinstance(result, list):
        for portal_data in result:
            for key, articles in portal_data.items():
                # Limit the number of items in the JSON list to n_value
                articles = articles[:n_value]

                if key in combined_results:
                    combined_results[key].extend(articles)
                else:
                    combined_results[key] = articles

    return combined_results


def save_combined_results(combined_results, successful_count, config):
    """Saves the combined results to JSON and text files."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # folder_dir = config['file_paths']['generated_article_links']
    folder_dir = config['file_paths']['generated_article_links']
    output_file_path = os.path.join(folder_dir, f"links_{timestamp}.json")

    # Ensure the directory exists
    os.makedirs(folder_dir, exist_ok=True)

    # Save combined results to JSON file
    with open(output_file_path, 'w') as f:
        json.dump(combined_results, f, indent=4)
    st.info(f"Successfully saved links from {successful_count} portal(s) to file: {output_file_path}")

    # Save links to edited_link_file
    save_links_to_txt(combined_results, config)


def save_links_to_txt(combined_results, config):
    """
    Saves the article links to the edited_link_file from config.

    Args:
        combined_results (dict): Dictionary containing search results
        config (dict): Configuration dictionary containing file paths
    """
    try:
        # Get file path from config
        if 'file_paths' not in config or 'edited_link_file' not in config['file_paths']:
            raise KeyError("Missing required config paths: file_paths.edited_link_file")

        filepath = config['file_paths']['edited_link_file']

        # Ensure directory exists
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        # Read existing links to avoid duplicates
        existing_links = set()
        if os.path.exists(filepath):
            with open(filepath, 'r') as f:
                existing_links = set(f.read().splitlines())

        # Collect all new valid URLs
        new_urls = set()

        # Handle different data structures in combined_results
        if isinstance(combined_results, dict):
            if 'search_results' in combined_results:
                # Handle search results format
                for item in combined_results['search_results']:
                    if isinstance(item, dict):
                        url = item.get('url')
                    else:
                        url = item
                    if url and isinstance(url, str) and is_valid_url(url):
                        new_urls.add(url)
            else:
                # Handle portal data format
                for articles in combined_results.values():
                    for article in articles:
                        if isinstance(article, dict):
                            url = article.get('url')
                        else:
                            url = article
                        if url and isinstance(url, str) and is_valid_url(url):
                            new_urls.add(url)

        # Add only new URLs that don't exist in the file
        new_urls = new_urls - existing_links

        # Append new URLs to the file
        if new_urls:
            with open(filepath, 'a') as f:
                for url in new_urls:
                    f.write(f"{url}\n")

            st.success(f"Added {len(new_urls)} new links to {filepath}")
            logger.info(f"Successfully added {len(new_urls)} new links to {filepath}")
        else:
            st.info("No new links to add (all URLs already exist in the file)")
            logger.info("No new links added - all URLs already exist in the file")

    except KeyError as e:
        error_msg = f"Configuration error: {str(e)}"
        logger.error(error_msg)
        st.error(error_msg)
    except Exception as e:
        error_msg = f"Error saving links to file: {str(e)}"
        logger.error(error_msg, exc_info=True)
        st.error(error_msg)


# def manage_link_content():
#     """Handles the UI for managing link content."""
#     with st.expander("📌 Manage Link Generation", expanded=True):
#         tabs = st.tabs(["🔗 Manual Link Input", "🤖 Automated Link Generator", "🔍 Add Links from Keywords"])  # New tab added

#         with tabs[0]:
#             link_content_manager()  # Call the link_content_manager function

#         with tabs[1]:
#             fetch_ai_news_links()

#         with tabs[2]:  # New tab for adding links from keywords
#             add_links_from_keywords()

def manage_link_content():
    """Enhanced UI for managing link content while maintaining all functionality."""

    st.markdown("""
        <style>
        /* Enhanced Tab Styling */
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
            padding: 8px;
            background-color: #f8f9fa;
            border-radius: 10px;
        }

        .stTabs [data-baseweb="tab"] {
            padding: 8px 16px;
            border-radius: 6px;
            font-weight: 500;
            background-color: transparent;
            border: 1px solid #dee2e6;
            transition: all 0.2s ease;
        }

        .stTabs [aria-selected="true"] {
            background-color: #007bff !important;
            color: white !important;
            border-color: #007bff !important;
            border-radius: 6px !important;
        }

        /* Button Enhancements */
        .stButton button {
            width: 100%;
            margin-top: 10px;
            height: 45px;
            border-radius: 6px;
            font-weight: 500;
            transition: all 0.2s ease;
            border: 2px solid #007bff;
        }

        .stButton button:hover {
            transform: translateY(-1px);
            box-shadow: 0 2px 6px rgba(0,0,0,0.1);
        }

        /* Input Field Styling */
        .stTextInput input, .stTextArea textarea {
            border: 2px solid #e9ecef;
            border-radius: 6px;
            padding: 12px;
            transition: all 0.2s ease;
        }

        .stTextInput input:focus, .stTextArea textarea:focus {
            border-color: #007bff;
            box-shadow: 0 0 0 2px rgba(0,123,255,0.25);
        }

        /* Info Box Styling */
        .info-box {
            background-color: #f8f9fa;
            border-left: 4px solid #007bff;
            padding: 12px 16px;
            margin: 10px 0;
            border-radius: 4px;
        }

        /* Helper Text Styling */
        .helper-text {
            color: #6c757d;
            font-size: 0.9em;
            margin-bottom: 16px;
            padding: 12px;
            background-color: #f8f9fa;
            border-radius: 6px;
            border: 1px solid #e9ecef;
        }

        /* Section Headers */
        h3 {
            color: #2c3e50;
            margin-bottom: 16px;
            font-weight: 600;
        }

        /* Alert/Info Message Styling */
        .stAlert {
            padding: 12px 16px;
            border-radius: 6px;
            margin: 10px 0;
        }

        /* Progress Bar Enhancement */
        .stProgress > div > div {
            background-color: #007bff;
        }

        /* Expander Styling */
        .streamlit-expanderHeader {
            background-color: #f8f9fa;
            border-radius: 6px;
            padding: 8px 12px;
            font-weight: 500;
        }

        /* Metrics Styling */
        [data-testid="stMetricValue"] {
            font-size: 1.2rem;
            font-weight: 600;
            color: #007bff;
        }
        </style>
    """, unsafe_allow_html=True)

    with st.expander("📌 Manage Link Generation", expanded=True):
        # Enhanced top-level helper text
        st.markdown("""
            <div class='helper-text'>
            <b>Choose your preferred method:</b><br>
            🔸 <b>Manual:</b> Directly paste your links<br>
            🔸 <b>Automated:</b> Generate links from news sources<br>
            🔸 <b>Keywords:</b> Find links based on search terms
            </div>
        """, unsafe_allow_html=True)

        tabs = st.tabs([
            "🔗 Manual Link Input",
            "🤖 Automated Link Generator",
            "🔍 Add Links from Keywords"
        ])

        with tabs[0]:
            st.markdown("### Manual Link Input")
            st.markdown("""
                <div class='info-box'>
                📝 <b>Instructions:</b><br>
                • Paste your links below (one per line)<br>
                • Links will be validated automatically<br>
                • Supported: News articles, blog posts, and web pages
                </div>
            """, unsafe_allow_html=True)

            link_content_manager()

        with tabs[1]:
            st.markdown("### Automated Link Generator")
            st.markdown("""
                <div class='info-box'>
                🤖 <b>How it works:</b><br>
                1. Enter source URLs below<br>
                2. Set the number of links to generate<br>
                3. Click Generate to start the process
                </div>
            """, unsafe_allow_html=True)

            fetch_ai_news_links()

        with tabs[2]:
            st.markdown("### Keyword-Based Search")
            st.markdown("""
                <div class='info-box'>
                🔍 <b>Search options:</b><br>
                • Enter keywords separated by commas<br>
                • Adjust the maximum results count<br>
                • Select your preferred time range
                </div>
            """, unsafe_allow_html=True)

            config = load_config()
            if not config:
                st.error("⚠️ Configuration loading failed")
                return

            col1, col2 = st.columns([2, 1])

            with col1:
                keywords = st.text_input(
                    "Keywords (comma-separated)",
                    help="Example: AI technology, machine learning, data science"
                )

            with col2:
                max_links = st.number_input(
                    "Max results",
                    min_value=1,
                    max_value=50,
                    value=10,
                    help="Maximum number of links to retrieve"
                )

            time_filter = st.select_slider(
                "Time range",
                options=['d', 'w', 'm', 'y'],
                value='w',
                format_func=lambda x: {
                    'd': 'Last 24 hours',
                    'w': 'Past week',
                    'm': 'Past month',
                    'y': 'Past year'
                }[x]
            )

            if st.button("🔎 Search and Add Links", use_container_width=True):
                if not keywords.strip():
                    st.warning("⚠️ Please enter at least one keyword")
                    return

                with st.spinner("🔍 Searching for relevant links..."):
                    process_keyword_search(keywords, max_links, time_filter, config)

def process_keyword_search(keywords, max_links, time_filter, config):
    """Enhanced keyword search processing with better feedback"""
    try:
        search_module = SearchModule()
        keyword_list = [k.strip() for k in keywords.split(',')]

        progress_bar = st.progress(0)
        status_text = st.empty()

        all_links = []
        for idx, keyword in enumerate(keyword_list):
            status_text.text(f"Searching for: {keyword}")
            progress_bar.progress((idx + 1) / len(keyword_list))

            results = search_module.search_duckduckgo(
                keywords=keyword,
                timelimit=time_filter,
                max_results=max_links
            )

            if results:
                relevant_links = search_module.filter_relevant_links(
                    search_results=results,
                    keywords=[keyword]
                )
                all_links.extend(relevant_links)

        # Remove duplicates while preserving order
        all_links = list(dict.fromkeys(all_links))[:max_links]

        if all_links:
            save_links_to_txt({'search_results': [{'url': url} for url in all_links]}, config)

            # Show results summary
            st.success(f"✅ Successfully added {len(all_links)} new links!")

            # Display metrics in a more visual way
            metrics = search_module.get_performance_metrics()
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Search Time", f"{metrics['search_duration']:.1f}s")
            with col2:
                st.metric("Filter Time", f"{metrics['filter_duration']:.1f}s")
            with col3:
                st.metric("Relevance", f"{metrics['relevance_rate']}%")

            # Show found links in an expandable section
            with st.expander("📋 View Found Links"):
                for link in all_links:
                    st.write(f"• {link}")
        else:
            st.warning("No relevant links found matching your criteria")

    except Exception as e:
        st.error(f"Error during search: {str(e)}")
        logger.error(f"Search error: {str(e)}", exc_info=True)

def add_links_from_keywords():
    """Handles the UI for adding links based on user-defined keywords."""
    st.header("🔍 Add Links from Keywords")

    # Load configuration
    config = load_config()
    if not config:
        st.error("Failed to load configuration")
        return

    keywords = st.text_input("Enter keywords (comma-separated):", help="Keywords to search for links.")
    max_links = st.number_input("Max number of links to extract:", min_value=1, value=10, help="Set the maximum number of links to save.")
    time_filter = st.selectbox(
        "Time filter:",
        options=['d', 'w', 'm', 'y'],
        index=1,
        format_func=lambda x: {
            'd': 'Past 24 hours',
            'w': 'Past week',
            'm': 'Past month',
            'y': 'Past year'
        }[x],
        help="Filter results by time"
    )

    if st.button("Search and Add Links"):
        if keywords:
            # Initialize SearchModule
            search_module = SearchModule()
            keyword_list = [keyword.strip() for keyword in keywords.split(',')]
            all_links = []

            with st.spinner("Searching for links..."):
                for keyword in keyword_list:
                    # Use the SearchModule's search_duckduckgo method
                    search_results = search_module.search_duckduckgo(
                        keywords=keyword,
                        timelimit=time_filter,
                        max_results=max_links
                    )

                    if search_results:
                        # Filter relevant links using the module's method
                        relevant_links = search_module.filter_relevant_links(
                            search_results=search_results,
                            keywords=[keyword]
                        )
                        all_links.extend(relevant_links)

                # Remove duplicates while preserving order
                all_links = list(dict.fromkeys(all_links))

                # Limit the number of links to the user-defined maximum
                all_links = all_links[:max_links]

                if all_links:
                    # Save the links using the configured file path
                    save_links_to_txt({'search_results': [{'url': url} for url in all_links]}, config)
                    st.success(f"Added {len(all_links)} new links!")

                    # Display performance metrics
                    metrics = search_module.get_performance_metrics()
                    st.info(f"""
                    Search Performance:
                    - Search Duration: {metrics['search_duration']:.2f} seconds
                    - Filter Duration: {metrics['filter_duration']:.2f} seconds
                    - Relevance Rate: {metrics['relevance_rate']}%
                    """)
                else:
                    st.warning("No relevant links found that meet the criteria.")
        else:
            st.warning("Please enter at least one keyword.")


def save_links_to_txt(combined_results, config):
    """
    Saves the article links to a text file using the configured file path.

    Args:
        combined_results (dict): Dictionary containing search results
        config (dict): Configuration dictionary containing file paths
    """
    try:
        filepath = config['file_paths']['edited_link_file']

        # Ensure directory exists
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        # Read existing links to avoid duplicates
        existing_links = set()
        if os.path.exists(filepath):
            with open(filepath, 'r') as f:
                existing_links = set(f.read().splitlines())

        # Add new links
        new_links_count = 0
        with open(filepath, 'a') as f:
            for portal_data in combined_results.values():
                for article in portal_data:
                    url = article['url']
                    if url not in existing_links and is_valid_url(url):
                        f.write(f"{url}\n")
                        existing_links.add(url)
                        new_links_count += 1

        st.info(f"Added {new_links_count} new links to the configured file: {filepath}")

        # Log the operation
        logger.info(f"Successfully added {new_links_count} new links to {filepath}")

    except KeyError as e:
        error_msg = f"Missing configuration key: {str(e)}"
        logger.error(error_msg)
        st.error(error_msg)
    except Exception as e:
        error_msg = f"Error saving links to file: {str(e)}"
        logger.error(error_msg, exc_info=True)
        st.error(error_msg)


def load_and_display_config():
    """Loads and displays the configuration."""
    try:
        config = load_config()
    except Exception as e:
        st.error(f"⚠️ Error loading configuration: {e}")
        logger.error(f"Error loading configuration: {e}", exc_info=True)
        return None
    return config


def process_links(config):
    """Handles the processing of articles based on user input."""
    if not config:
        return

    with st.expander("⚙️ Process Links", expanded=True):
        tabs = st.tabs(["🔍Review and Edit Prompt", "🚀 Process Articles"])

        with tabs[0]:
            prompt = manage_prompt(config)

        with tabs[1]:
            selected_urls = select_articles_to_process(config)
            if selected_urls:
                process_selected_articles(selected_urls, config, prompt)


def manage_prompt(config):
    """Handles prompt selection and editing."""
    show_prompt = st.checkbox("Show and edit prompt", help="Enable this to view and modify the processing prompt")
    if show_prompt:
        prompt = load_and_edit_prompt(config)
        if prompt is None:
            st.warning("⚠️ No prompt selected or available.")
    else:
        prompt = load_main_prompt(config)
    return prompt



def select_articles_to_process(config):
    """Manages article selection for processing."""
    all_urls = display_links_selection(config)

    st.caption("Select whether to process specific articles or all available articles 🔄")
    process_option = st.radio("Choose processing option: 📈", ["Select specific articles", "Process all articles"])

    if process_option == "Select specific articles":
        selected_urls = st.multiselect("Select articles to process: 📚", all_urls, help="Choose one or more articles from the list to process")
    else:
        selected_urls = all_urls

    return selected_urls


def process_selected_articles(selected_urls, config, prompt):
    """Processes the selected articles and handles the results."""
    st.caption("Start processing the selected articles using the chosen prompt ⚙️")
    if st.button("🚀 Confirm and Process Selected Links"):
        with st.spinner("Processing selected links... ⏳"):
            try:
                results, total_duration = process_urls(selected_urls, config, prompt)

                # Handle results
                handle_processing_results(results, config, total_duration)

            except Exception as e:
                st.error(f"❌ Error during processing: {e}")
                logger.error(f"Error during processing: {e}", exc_info=True)


def handle_processing_results(results, config, total_duration):
    """
    Saves and displays the results of the processed articles.
    """
    # Filter out skipped articles for file saving
    results_to_save = [r for r in results if not r.get('skipped', False)]

    # Save results to file
    json_file_path = save_results_to_file(results_to_save, config)

    # Store results and file_path in session state
    st.session_state.results = results
    st.session_state.json_file_path = json_file_path
    st.session_state.total_duration = total_duration

    # Count skipped articles
    skipped_count = sum(1 for r in results if r.get('skipped', False))

    try:
        # Initialize MongoDB handler
        mongo_handler = SummaryMongoHandler()
        user_email = st.session_state.get('user_email')

        if user_email:
            # Save to MongoDB (only non-skipped results)
            valid_results = [r for r in results_to_save if 'error' not in r]

            if valid_results:
                save_result = mongo_handler.save_summaries_to_mongo(valid_results, user_email)
                logger.info(f"Save result: {save_result}")

                if save_result["success"]:
                    st.success(f"""✅ Results saved:
                    - Local JSON: {json_file_path}
                    - MongoDB: {save_result['saved_count']} new articles
                    - Skipped: {skipped_count} already existing articles
                    - Processing time: {total_duration:.2f} seconds""")
                else:
                    st.error(f"❌ MongoDB save failed: {save_result.get('error', 'Unknown error')}")
                    st.success(f"✅ Results saved to local JSON: {json_file_path}")
                    st.info(f"⏱ Processing time: {total_duration:.2f} seconds")
            else:
                st.warning(f"No new articles to save. {skipped_count} articles were already processed.")
        else:
            st.warning("⚠️ Results saved locally only. Log in to save to MongoDB.")
            st.success(f"✅ Results saved to local JSON: {json_file_path}")
            st.info(f"⏱ Processing time: {total_duration:.2f} seconds")

    except Exception as e:
        logger.error(f"MongoDB save error: {str(e)}", exc_info=True)
        st.error("❌ Error saving to MongoDB. Results saved locally only.")
        st.success(f"✅ Results saved to local JSON: {json_file_path}")
        st.info(f"⏱ Processing time: {total_duration:.2f} seconds")


def view_and_export_article_summaries():
    """Displays the summaries of processed articles."""
    with st.expander("📄 View Summaries", expanded=False):
        view_export_article_summaries()  # Call the function to display article summaries


def main():
    st.title("🔍 Smart Article-Link Summarizer")

    # Display welcome message with user email
    st.markdown(f"Welcome, **{user_email}**! 👋")

    # Link content management
    manage_link_content()

    # Configuration loading
    config = load_and_display_config()

    # Link processing management
    process_links(config)

    # View article summaries
    view_and_export_article_summaries()

if __name__ == "__main__":
    main()